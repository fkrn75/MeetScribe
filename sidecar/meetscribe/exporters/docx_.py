"""DOCX 내보내기 — Word 회의록 문서(화자별 문단).

python-docx 로 .docx 를 생성한다. 무거운(외부) 의존성이므로 **함수 안에서
lazy import** 한다 — 이 모듈을 import 하는 것만으로 python-docx 가 필요해지지
않게(미설치 환경에서 exporters 패키지 로드가 깨지지 않게).

레이아웃:
    제목(회의록) → 메타(언어·길이·화자 수) → 연속 발화를 화자별 문단으로.
연속한 동일 화자 세그먼트는 한 문단으로 합쳐 가독성을 높인다.
"""

from __future__ import annotations

from ..schemas import TranscriptionResult

_UNKNOWN_SPEAKER = "화자 미상"


def _fmt_timestamp(seconds: float) -> str:
    """초 → 'mm:ss'(60분 넘으면 분 누적)."""
    total = int(seconds) if seconds and seconds > 0 else 0
    minutes, secs = divmod(total, 60)
    return f"{minutes:02d}:{secs:02d}"


def write(result: TranscriptionResult, out_path: str) -> str:
    """결과를 Word 문서로 만들어 out_path 에 저장한다."""
    # lazy import: 미설치 환경에서도 모듈 로드는 통과(호출 시점에만 필요).
    from docx import Document  # type: ignore

    doc = Document()

    # ── 제목 + 메타 정보 ─────────────────────────────────────
    doc.add_heading("회의록", level=0)
    speaker_count = len(result.speakers) if result.speakers else 0
    meta = doc.add_paragraph()
    meta.add_run(
        f"언어: {result.language}    "
        f"길이: {_fmt_timestamp(result.duration)}    "
        f"화자 수: {speaker_count}"
    ).italic = True

    doc.add_paragraph()  # 메타와 본문 사이 여백

    # ── 본문: 연속 동일 화자 발화를 한 문단으로 묶는다 ──────────
    cur_speaker: str | None = None
    cur_start: float = 0.0
    buffer: list[str] = []

    def flush() -> None:
        """모아둔 발화를 한 문단으로 출력한다."""
        if not buffer:
            return
        speaker = cur_speaker or _UNKNOWN_SPEAKER
        para = doc.add_paragraph()
        # '[mm:ss] 화자:' 를 굵게, 이어서 합친 텍스트.
        para.add_run(f"[{_fmt_timestamp(cur_start)}] {speaker}: ").bold = True
        para.add_run(" ".join(buffer))

    for seg in result.segments:
        text = (seg.text or "").strip()
        if not text:
            continue
        if seg.speaker != cur_speaker:
            # 화자가 바뀌면 직전 묶음을 내보내고 새 묶음 시작.
            flush()
            cur_speaker = seg.speaker
            cur_start = seg.start
            buffer = [text]
        else:
            buffer.append(text)

    flush()  # 마지막 묶음

    doc.save(out_path)
    return out_path
